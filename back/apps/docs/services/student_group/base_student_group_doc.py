import datetime
import io
import os
import re
import uuid
from abc import abstractmethod

from django.db.models import QuerySet
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from openpyxl.workbook import Workbook
from pandas._libs.tslibs.offsets import BDay

from apps.applications.consts.application_statuses import DRAFT
from apps.applications.selectors.course_application import course_application_orm
from apps.applications.selectors.event_application import event_application_orm
from apps.commons.utils.django.settings import settings_utils
from apps.commons.utils.ldap import ldap_utils
from apps.docs.utils.docx_utils import generate_word_file
from apps.edu.exceptions.student_group.curator_not_set import CuratorNotSet
from apps.edu.exceptions.student_group.student_group_not_found import StudentGroupNotFound
from apps.edu.selectors.student_group import student_group_orm


class BaseStudentGroupDoc:
    """
    Базовый класс для формирования документов по учебной группе
    """

    student_group = None

    def __init__(self, group_id: uuid):
        """
        Инициализация класса - получение документа по учебной группе
        :param group_id: object_id учебной группы
        """
        self.student_group = student_group_orm.get_one_record_or_none({'object_id': group_id})
        if not self.student_group:
            raise StudentGroupNotFound
        if not self.student_group.curator:
            raise CuratorNotSet

    def _get_department_and_manager(self) -> tuple:
        """
        Получение подразделения и его менеджера
        :return:
        """
        if self.student_group.ou:
            dep = self.student_group.ou.program.department.display_name
        else:
            dep = self.student_group.iku.department.display_name
        all_centres = ldap_utils.get_departments_with_managers()
        manager = all_centres.get(dep, '')
        return dep, manager

    def _get_department_short_name(self) -> str:
        """
        Получение сокращения названия подразделения
        :return: Сокращение названия центра
        """
        if self.student_group.ou:
            department = self.student_group.ou.program.department.display_name
        else:
            department = self.student_group.iku.department.display_name
        dep_name_split = re.split(' |-', department)
        short_dep = ''
        for word in dep_name_split:
            short_dep += word[:1].upper()
        return short_dep

    def _get_date_start_and_end(self) -> tuple:
        """
        Получение даты начала и окончания обучения в учебной группе
        :return: кортеж дат
        """
        if self.student_group.ou:
            return self.student_group.ou.date_start, self.student_group.ou.date_end
        return self.student_group.iku.date_start, self.student_group.iku.date_end

    def _get_audience_categories_str(self) -> str:
        """
        Получение строкового представления для всех связанных с учебной группой
        категорий слушателей с разеделителем ","
        :return: строка с категориями слушателей через запятую
        """
        if self.student_group.ou:
            categories = self.student_group.ou.program.categories.all()
        else:
            categories = self.student_group.iku.categories.all()
        categories_string = ''
        for category in categories:
            categories_string += f'{category.name}, '
        return categories_string[:-2]

    @staticmethod
    def _get_order_date(date_start: datetime) -> datetime:
        """
        Получение даты размещения договора оферты
        :param date_start: дата начала обучения
        :return: дата
        """
        return date_start - BDay(settings_utils.get_parameter_from_settings('ORDER_DATE_DAYS'))

    @staticmethod
    def _get_pay_date(date_start: datetime) -> datetime:
        """
        Получение крайней даты оплаты
        :param date_start: дата начала обучения
        :return: дата
        """
        return date_start - BDay(settings_utils.get_parameter_from_settings('PAY_DATE_DAYS'))

    def _get_price(self) -> int:
        """
        Получение стоимости обучения
        :return: целое число, рублей
        """
        if self.student_group.ou:
            return self.student_group.ou.program.price
        return self.student_group.iku.price

    def _get_duration(self) -> int:
        """
        Получение количества часов - объем программы обучения
        :return: количество часов
        """
        if self.student_group.ou:
            return self.student_group.ou.program.duration
        return self.student_group.iku.duration

    def _get_application_count(self) -> int:
        """
        Получение количества заявок в учебной группе
        :return: целое число заявок
        """
        if self.student_group.ou:
            return (course_application_orm.
                    get_filter_records(filter_by={'group_id': self.student_group.object_id}).
                    count())
        return (event_application_orm.
                get_filter_records(filter_by={'group_id': self.student_group.object_id}).
                count())

    def _get_group_applications(self) -> QuerySet:
        """
        Получение QuerySet с заявками в учебную группу
        :return: QuerySet с завяками обучающихся
        """
        if self.student_group.ou:
            return (course_application_orm.get_filter_records(
                filter_by={'group_id': self.student_group.object_id},
                exclude={'status': DRAFT},
                order_by=['profile__surname', 'profile__name', 'profile__patronymic']
            ))
        return (event_application_orm.get_filter_records(
            filter_by={'group_id': self.student_group.object_id},
            exclude={'status': DRAFT},
            order_by=['profile__surname', 'profile__name', 'profile__patronymic']
        ))

    def _get_program_and_certificates_types(self) -> tuple:
        """
        Получение кортежа с типом ДПП и типом выдаваемого сертификата
        :return: кортеж из двух строк - тип ДПП, тип сертификата
        """
        if self.student_group.ou.program.type == 'Повышение квалификации':
            type_dpp = 'повышения квалификации'
            cert_type = 'повышении квалификации'
        else:
            type_dpp = 'профессиональной переподготовки'
            cert_type = 'профессиональной переподготовке'
        return type_dpp, cert_type

    def _get_template_path(self, dir_list: tuple) -> str:
        """
        Получение пути до шаблона файла
        :param dir_list: Кортеж с директориями до шаблона
        :return: Путь к шаблону
        """
        path = os.path.join(
            settings_utils.get_parameter_from_settings('MEDIA_ROOT'),
            'Шаблоны'
        )
        for folder in dir_list:
            path = os.path.join(path, folder)
        if self.student_group.ou:
            return str(os.path.join(path, 'шаблон_ОУ.docx'))
        return str(os.path.join(path, 'шаблон_ИКУ.docx'))

    @abstractmethod
    def _get_context(self) -> dict:
        """
        Получение словаря с данными для вставки
        :return: словарь с данными
        """
        pass

    def _generate_wb(self) -> Workbook():
        """
        Генерация эксель файла
        :return:
        """
        wb = Workbook()
        self._fill_wb(wb)
        return wb

    def _fill_wb(self, wb: Workbook):
        """
        Заполнение эксель файла
        :param wb: Эксель файл для заполнения
        :return:
        """
        pass

    def _set_student_data_into_table(self, doc, table_index: int = 0):
        """
        Подстановка данных о студентах в первую таблицу полученного файла ворд
        :param doc: Word файл
        :param table_index: Номер таблицы в документе (по умолчанию первая таблица)
        :return:
        """
        file_stream = io.BytesIO()
        doc.save(file_stream)
        document = Document(file_stream)
        table = document.tables[table_index]
        applications = self._get_group_applications()
        for index, application in enumerate(applications):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[1].text = (f'{application.profile.surname} '
                                 f'{application.profile.name} '
                                 f'{application.profile.patronymic}')
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(12)
        return document

    def get_response(self, folder_tuple: tuple, xlsx: bool = True) -> HttpResponse:
        """
        Получение HTTP респонза с файлом
        :param folder_tuple: Кортеж с директориями после "Шаблоны" до файла шаблона
        :param xlsx: Булево, если True - выгрузка эксель, иначе ворд
        :return:
        """
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        if xlsx:
            content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response = HttpResponse(content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="info_{self.student_group.code}.docx"'
        if xlsx:
            doc = self._generate_wb()
        else:
            doc = generate_word_file(
                self._get_template_path(folder_tuple),
                self._get_context()
            )
        doc.save(response)
        return response
