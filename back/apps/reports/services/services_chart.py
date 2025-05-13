import os

from django.db.models import QuerySet, Q
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from apps.commons.utils.django.settings import settings_utils
from apps.edu.selectors.services.education_service import education_service_orm
from apps.edu.selectors.services.information_service import information_service_orm
from apps.edu.selectors.student_group import student_group_model


# Путь к шаблону отчета графика услуг
service_chart_path = os.path.join(
    settings_utils.get_parameter_from_settings('MEDIA_ROOT'),
    'Шаблоны',
    'Отчеты',
    'шаблон_график_услуг.docx'
)


class ServiceChart:
    """Класс для формирования графика оказанных услуг за период"""

    def __init__(self, report_data: dict):
        """
        Инициализация класса - фиксация года и месяца отчета
        :param report_data: параметры отчета, содержащие год и месяц
        """
        self.report_year = report_data.get('report_year')
        self.report_month = report_data.get('report_month')

    def _get_student_groups(self) -> QuerySet:
        """
        Получение queryset со списком учебных групп попадающих под параметры отчета
        :return: qs
        """
        edu_services = education_service_orm.get_filter_records(
            filter_by=dict(date_start__year=int(self.report_year))
        )
        info_services = information_service_orm.get_filter_records(
            filter_by=dict(date_start__year=int(self.report_year))
        )
        if self.report_month != 'all':
            edu_services = edu_services.filter(date_start__month=int(self.report_month))
            info_services = info_services.filter(date_start__month=int(self.report_month))
        edu_ids = [service.object_id for service in edu_services]
        info_ids = [service.object_id for service in info_services]
        return student_group_model.objects.filter(Q(ou_id__in=edu_ids) | Q(iku_id__in=info_ids)).order_by('-object_id')

    def generate_file(self) -> Document:
        """Генерация word файла с графиком оказанных услуг попадающих под параметры отчета"""
        # Открыть шаблон файла отчета ДПП
        document = Document(service_chart_path)
        # Получение таблицы из файла отчета
        table = document.tables[0]
        groups = self._get_student_groups()
        # Заполнение таблицы
        for index, group in enumerate(groups, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            if group.iku is None:
                row_cells[1].text = group.ou.program.department.display_name
                if group.ou.program.type == 'Повышение квалификации':
                    row_cells[2].text = 'Курс повышения квалификации'
                else:
                    row_cells[2].text = 'Курс профессиональной переподготовки'
                row_cells[3].text = group.ou.program.name
                row_cells[4].text = str(group.ou.program.duration)
                row_cells[5].text = group.ou.date_start.strftime('%d.%m.%Y')
                row_cells[6].text = group.ou.date_end.strftime('%d.%m.%Y')
            else:
                row_cells[1].text = group.iku.department.display_name
                row_cells[2].text = group.iku.type.name
                row_cells[3].text = group.iku.name
                row_cells[4].text = str(group.iku.duration)
                row_cells[5].text = group.iku.date_start.strftime('%d.%m.%Y')
                row_cells[6].text = group.iku.date_end.strftime('%d.%m.%Y')
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(11)
        return document
