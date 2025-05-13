import os

from django.db.models import QuerySet
from docx import Document
from docx.shared import Pt

from apps.commons.utils.django.settings import settings_utils
from apps.edu.selectors.program import program_orm, program_model
from apps.edu.selectors.services.education_service import education_service_orm

# Путь к шаблону отчета ДПП
dpp_path = os.path.join(
    settings_utils.get_parameter_from_settings('MEDIA_ROOT'),
    'Шаблоны',
    'Отчеты',
    'шаблон_ДПП.docx'
)


class DppReport:
    """Класс для формирования списка ДПП за период"""

    def __init__(self, report_data: dict):
        """
        Инициализация класса - фиксация года и месяца отчета
        :param report_data: параметры отчета, содержащие год и месяц
        """
        self.report_year = report_data.get('report_year')
        self.report_month = report_data.get('report_month')

    def _get_program_queryset(self) -> QuerySet:
        """
        Получение queryset со списком программ попадающим под параметры запроса
        :return: qs
        """
        courses = education_service_orm.get_filter_records(
            filter_by=dict(date_start__year=int(self.report_year))
        )
        if self.report_month != 'all':
            courses = courses.filter(date_start__month=int(self.report_month))
        dpp_ids = [course.program_id for course in courses]
        return program_orm.get_filter_records(filter_by=dict(object_id__in=dpp_ids), order_by=['-object_id'])

    @staticmethod
    def _get_program_categories_text(program: program_model) -> str:
        """
        Получение строкового представления всех категорий слушателей программы
        :param program: ДПП
        :return: текст со всеми категориями слушателей ДПП
        """
        text = ''
        for cat in program.categories.all():
            text += f'{cat.name};\n'
        return text

    def generate_file(self) -> Document:
        """Генерация word файла со списком ДПП попадающим под параметры отчета"""
        # Открыть шаблон файла отчета ДПП
        document = Document(dpp_path)
        # Получение таблицы из файла отчета
        table = document.tables[0]
        programs = self._get_program_queryset()
        # Заполнение таблицы
        for index, dpp in enumerate(programs, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = dpp.name
            row_cells[2].text = str(dpp.duration)
            row_cells[3].text = self._get_program_categories_text(dpp)
            row_cells[4].text = dpp.annotation
            if dpp.program_order is None:
                row_cells[5].text = '-'
            else:
                row_cells[5].text = f'№{dpp.program_order.number} от {dpp.program_order.date.strftime("%d.%m.%Y")}'
        # Форматирование всех ячеек таблицы
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(11)
        return document
